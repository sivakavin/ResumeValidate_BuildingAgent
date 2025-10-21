
import streamlit as st
import tempfile
import os
# Add project root to sys.path for module imports
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from graph.resume_flow import ResumeFlow
from dotenv import load_dotenv


import docx2txt
from PyPDF2 import PdfReader
import io
from docx import Document
import re

def extract_text_from_file(uploaded_file):
    if uploaded_file.type == "application/pdf":
        reader = PdfReader(uploaded_file)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text
    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name
        text = docx2txt.process(tmp_path)
        os.unlink(tmp_path)
        return text
    else:
        return uploaded_file.read().decode("utf-8")

# Helper to highlight improvements (simple diff)
def highlight_changes(original, refined):
    orig_lines = set(original.splitlines())
    refined_lines = set(refined.splitlines())
    added = refined_lines - orig_lines
    highlighted = []
    for line in refined.splitlines():
        if line.strip() in added and line.strip():
            highlighted.append(f":green[**{line}**]")
        else:
            highlighted.append(line)
    return "\n".join(highlighted)

# Wrapper for processing resume
def process_resume(jd_text, resume_text):
    result = flow.run(jd_text, resume_text)
    supervisor_feedback = result["evaluation"].get("feedback", "No feedback.")
    refined_resume = result["resume"]
    return supervisor_feedback, refined_resume

# Helper to create DOCX for download
def create_docx(text):
    doc = Document()
    # Try to split and add headings for Skills, Experience, Education
    sections = {"Skills": [], "Experience": [], "Education": [], "Other": []}
    current = "Other"
    for line in text.splitlines():
        l = line.strip()
        if re.match(r"skills?[:]?", l, re.I):
            current = "Skills"
            continue
        elif re.match(r"experience[:]?", l, re.I):
            current = "Experience"
            continue
        elif re.match(r"education[:]?", l, re.I):
            current = "Education"
            continue
        if l:
            sections[current].append(l)
    # Add name/title if present
    lines = text.splitlines()
    if lines and lines[0].strip():
        doc.add_heading(lines[0].strip(), 0)
    # Add sections with headings
    for sec in ["Skills", "Experience", "Education", "Other"]:
        if sections[sec]:
            doc.add_heading(sec, level=1)
            for item in sections[sec]:
                doc.add_paragraph(item)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_pdf(text):
    """Create a simple PDF from plain text using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except Exception:
        raise RuntimeError("reportlab is required to generate PDF. Install with: pip install reportlab")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    margin = 72
    y = height - margin
    max_width = width - 2 * margin
    # simple line wrap
    from reportlab.pdfbase.pdfmetrics import stringWidth
    lines = []
    for paragraph in text.splitlines():
        if not paragraph:
            lines.append("")
            continue
        words = paragraph.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if stringWidth(test, "Helvetica", 10) < max_width:
                line = test
            else:
                lines.append(line)
                line = w
        if line:
            lines.append(line)

    for line in lines:
        if y < margin:
            c.showPage()
            y = height - margin
        c.setFont("Helvetica", 10)
        c.drawString(margin, y, line)
        y -= 12

    c.save()
    buf.seek(0)
    return buf

# Load environment variables
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")
flow = ResumeFlow(groq_api_key)

st.set_page_config(page_title="Resume Multi-Agent", layout="wide")

# Gentle theme / peaceful colors
st.markdown(
    """
    <style>
    .stApp { background: linear-gradient(180deg, #f7fbfc 0%, #f1f7f8 100%); }
    .sidebar .sidebar-content { background: #ffffffcc; border-radius: 8px; padding: 12px }
    .card { background: #ffffff; padding: 12px; border-radius: 8px; box-shadow: 0 1px 6px rgba(16,24,40,0.06); }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("Resume Analyser & Builder")
st.subheader("Smart, fast resume analysis and tailored rewrites")


if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "refined_resume" not in st.session_state:
    st.session_state.refined_resume = None
if "original_resume" not in st.session_state:
    st.session_state.original_resume = None

st.markdown("""
Upload your resume and paste the job description. Click 'Analyze and Refine Resume' to get feedback and a tailored resume.
""")

# Sidebar: inputs
with st.sidebar.expander("Inputs", expanded=True):
    uploaded_file = st.file_uploader("Upload Resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
    jd_text = st.text_area("Paste Job Description", height=200)
    export_same_format = st.checkbox("Export in same format as uploaded (if supported)", value=True)
    max_iterations = st.number_input("Max refinement iterations", min_value=1, max_value=5, value=3)
    st.markdown("\n")
    st.markdown("Need a quick tip: upload a clean resume PDF or DOCX for best results.")



# Enhanced process_resume for multiple iterations and chat turns
def process_resume(jd_text, resume_text, max_loops=3):
    chat_turns = []
    current_resume = resume_text
    for i in range(max_loops):
        result = flow.run(jd_text, current_resume, max_loops=1)
        supervisor_feedback = result["evaluation"].get("feedback", "No feedback.")
        refined_resume = result["resume"]
        chat_turns.append({
            "role": "user",
            "content": f"**Job Description:**\n{jd_text}\n\n**Resume:**\n{current_resume}",
            "align": "left"
        })
        chat_turns.append({
            "role": "supervisor",
            "content": f"**Supervisor Feedback:**\n{supervisor_feedback}",
            "align": "right"
        })
        if not result["evaluation"].get("request_rebuild", False):
            break
        current_resume = refined_resume
    return chat_turns, refined_resume, resume_text

if st.button("Analyze and Refine Resume"):
    if not uploaded_file or not jd_text.strip():
        st.warning("Please upload a resume and paste a job description.")
    else:
        # Reset chat history for a new analysis
        st.session_state.chat_history = []
        with st.spinner("Analyzing and refining resume..."):
            resume_text = extract_text_from_file(uploaded_file)
            # Preliminary JD analysis for preview / scoring
            try:
                jd_analysis = flow.jd_analyzer.analyze(jd_text)
            except Exception:
                jd_analysis = {"skills": [], "tone": "", "keywords": []}

            # compute simple scores
            def compute_scores(jd_analysis, resume_text):
                skills = jd_analysis.get("skills", []) or []
                keywords = jd_analysis.get("keywords", []) or []
                rt = (resume_text or "").lower()
                skill_matches = sum(1 for s in skills if s.lower() in rt)
                keyword_matches = sum(1 for k in keywords if k.lower() in rt)
                skill_score = round((skill_matches / len(skills) * 5) if skills else 0, 1)
                keyword_score = round((keyword_matches / len(keywords) * 5) if keywords else 0, 1)
                tone = jd_analysis.get("tone", "").lower()
                tone_score = 5.0 if tone and tone in rt else 3.0
                # normalize
                def clamp(x):
                    if x is None:
                        return 0.0
                    return max(0.0, min(5.0, x))
                return {"skill_score": clamp(skill_score), "keyword_score": clamp(keyword_score), "tone_score": clamp(tone_score)}

            scores = compute_scores(jd_analysis, resume_text)

            # prepare improvisation ideas
            missing_skills = [s for s in (jd_analysis.get("skills") or []) if s.lower() not in (resume_text or "").lower()]
            improv_ideas = []
            if missing_skills:
                improv_ideas.append(f"Consider adding or emphasizing these skills: {', '.join(missing_skills[:8])}")
            if jd_analysis.get("keywords"):
                missing_kw = [k for k in jd_analysis.get("keywords") if k.lower() not in (resume_text or "").lower()]
                if missing_kw:
                    improv_ideas.append(f"Include important keywords: {', '.join(missing_kw[:8])}")

            # Run iterative refinement
            chat_turns, refined_resume, original_resume = process_resume(jd_text, resume_text, max_loops=max_iterations)
            st.session_state.original_resume = original_resume
            st.session_state.refined_resume = refined_resume
            st.session_state.chat_history.extend(chat_turns)

            # store analysis preview in session for right column
            st.session_state.jd_analysis = jd_analysis
            st.session_state.scores = scores
            st.session_state.improv_ideas = improv_ideas
            st.session_state.uploaded_type = uploaded_file.type




# Show only Supervisor Feedback (right-aligned) in chat window
for msg in st.session_state.chat_history:
    if msg.get("align") == "right":
        with st.container():
            st.markdown(f"<div style='text-align: right; background: #e6f7ff; padding: 10px; border-radius: 10px; margin: 5px 0;'>{msg['content']}</div>", unsafe_allow_html=True)

# Feedback & Analysis section (appears under Supervisor feedback)
jd_analysis = st.session_state.get("jd_analysis")
scores = st.session_state.get("scores")
improv_ideas = st.session_state.get("improv_ideas")
if jd_analysis or scores or improv_ideas:
    st.markdown("---")
    st.header("Feedback & Analysis")
    # JD analysis
    if jd_analysis:
        st.subheader("JD Analysis")
        try:
            st.json(jd_analysis)
        except Exception:
            st.write(str(jd_analysis))
    # Scores
    if scores:
        st.subheader("Analysis Scores (out of 5)")
        try:
            keys = list(scores.keys())
            cols = st.columns(len(keys)) if keys else [st]
            for i, k in enumerate(keys):
                cols[i].metric(k.replace("_", " ").title(), scores[k])
        except Exception:
            st.write(scores)
    # Improvisation ideas
    if improv_ideas:
        st.subheader("Improvisation Ideas")
        for idea in improv_ideas:
            st.info(idea)


# Show download button for final refined resume only
if st.session_state.refined_resume:
    st.markdown("---")
    # Decide output format
    out_buf = None
    out_filename = "refined_resume.docx"
    out_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    try:
        prefer_same = export_same_format
    except NameError:
        prefer_same = True

    uploaded_type = st.session_state.get("uploaded_type", None)
    if prefer_same and uploaded_type:
        if "pdf" in uploaded_type:
            try:
                out_buf = create_pdf(st.session_state.refined_resume)
                out_filename = "refined_resume.pdf"
                out_mime = "application/pdf"
            except Exception as e:
                st.warning(f"PDF generation failed: {e}. Falling back to DOCX.")
                out_buf = create_docx(st.session_state.refined_resume)
        elif "word" in uploaded_type or "officedocument.wordprocessingml" in uploaded_type or "docx" in uploaded_type:
            out_buf = create_docx(st.session_state.refined_resume)
            out_filename = "refined_resume.docx"
            out_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif "text" in uploaded_type or "plain" in uploaded_type:
            out_buf = io.BytesIO(st.session_state.refined_resume.encode("utf-8"))
            out_filename = "refined_resume.txt"
            out_mime = "text/plain"
        else:
            out_buf = create_docx(st.session_state.refined_resume)
    else:
        # default to DOCX
        out_buf = create_docx(st.session_state.refined_resume)

    if out_buf is not None:
        st.download_button(
            label=f"Download Refined Resume ({out_filename.split('.')[-1].upper()})",
            data=out_buf,
            file_name=out_filename,
            mime=out_mime,
        )
